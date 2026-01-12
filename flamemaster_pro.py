#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
FlameMaster Pro: A Comprehensive Multi-Platform Dynamic Analysis Tool
Author: Haroon Awan
Usage: FlameMaster_pro.py <application_file> [--custom "PatternName:Regex"]...
Description: FlameMaster Pro is an advanced cybersecurity tool designed to perform in-depth analysis of various file types,
including APKs, IPA, Windows executables, and more. It integrates both static and dynamic analysis techniques to extract
comprehensive information from malware samples, aiding in threat detection and remediation.
"""

import os
import sys
import zipfile
import re
import binascii
import datetime
import logging
import subprocess
import shutil
import xml.etree.ElementTree as ET
from pathlib import Path
from argparse import ArgumentParser, RawTextHelpFormatter

# External Libraries - Check if they are installed
required_packages = [
    'androguard',
    'rich',
    'colorama',
    'jinja2',
    'scapy',
    'pdfminer.six',
    'python-docx',
    'lief'
]

missing_packages = []
for package in required_packages:
    try:
        __import__(package.replace('-', '_').split('.')[0])
    except ImportError:
        missing_packages.append(package)

if missing_packages:
    print(f"Missing required packages: {', '.join(missing_packages)}")
    print("Install them with: pip install " + " ".join(missing_packages))
    sys.exit(1)

# Now import the packages
from androguard.misc import AnalyzeAPK
from androguard.core.bytecodes.apk import APK
from rich.console import Console
from rich.table import Table
from rich.panel import Panel
from rich.text import Text
from rich.progress import Progress, SpinnerColumn, BarColumn, TextColumn
from colorama import init as colorama_init
from jinja2 import Environment, FileSystemLoader
from scapy.all import rdpcap
from pdfminer.high_level import extract_text
from docx import Document
import lief

# Initialize colorama
colorama_init()

# Initialize Rich console
console = Console()

# Setup logging
log_dir = Path.cwd() / "logs"
log_dir.mkdir(exist_ok=True)
log_file = log_dir / f"FlameMaster_pro_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.log"

logging.basicConfig(
    filename=log_file,
    filemode='w',
    format='%(asctime)s - %(levelname)s - %(message)s',
    level=logging.DEBUG
)

# Color Codes (Using Rich for better colors)
SUCCESS = "[bold green]"
ERROR = "[bold red]"
WARNING = "[bold yellow]"
INFO = "[bold blue]"
RESET = "[/bold]"

# Define patterns to search (Expanded to cover extensive metadata extraction)
PATTERNS = {
    "Device ID References": r'getDeviceId',
    "Intent References": r'android\.intent',
    "Command Execution References": r'Runtime\.getRuntime\(\)\.exec',
    "SQLite References": r'SQLiteDatabase',
    "Logging References": r'Log\.(d|i|w|e)',
    "Content Providers": r'content://',
    "Broadcast Send References": r'sendBroadcast',
    "Service References": r'stopService|startService',
    "File References": r'file://',
    "SharedPreferences References": r'getSharedPreferences',
    "External Storage References": r'getExternal',
    "Crypto References": r'crypto\.',
    "MessageDigest References": r'MessageDigest',
    "Random Number Generators": r'java\.util\.Random',
    "Base64 References": r'Base64',
    "Hex References": r'Hex|hex\.',
    "Hardcoded Secrets": r'secret|password|username',
    "URL References": r'(www|http:|https:)',
    "HTTP Headers": r'addHeader',
    "Socket References": r'\.connect\(|\.disconnect|serverSocket|DatagramSocket',
    "SSL Certificate Files": r'\.(pkcs|p12|cer|der)',
    "SSL Certificate Pinning": r'getCertificatePinningSSL',
    "SSL Connections": r'ssl\.SSL',
    "WebView References": r'WebView',
    "JavaScript Interface": r'addJavascriptInterface',
    "JavaScript Enabled": r'setJavaScriptEnabled',
    "Domain References": r'domain',
    "Bucket References": r'buckets',
    "Storage References": r'storage',
    "API References": r'\bapi[_\-]?[\w]+\b',
    "Reporting References": r'report|reporting',
    "FTP References": r'ftp',
    "Endpoint References": r'endpoint|Endpoint|EndPoint',
    "Gopher References": r'gopher',
    "Mail References": r'smtp|mail|pop3|exchange',
    "HTTP References": r'\bhttp\b',
    "SQL Database References": r'(SELECT|INSERT|UPDATE|DELETE)\s+.*\s+(FROM|INTO|WHERE)',
    "Subdomain References": r'subdomain',
    "Email Addresses": r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,6}\b',
    "Google References": r'google',
    "Firebase References": r'firebaseio',
    "Facebook References": r'facebook',
    "Alibaba References": r'alibaba',
    "SSH References": r'ssh',
    "Dump References": r'dump',
    "API Variable References": r'_api|api_|_api_',
    "Valid IP Addresses": r'\b(?:[0-9]{1,3}\.){3}[0-9]{1,3}\b',
    "Obfuscation Patterns": r'^\s*(?:obfuscated|encrypt|decrypt)\s*',
    "Anti-Debugging Techniques": r'IsDebuggerPresent|checkDebugger',
    "Root Detection": r'isRooted|RootUtils',
    "Emulator Detection": r'emulator|Genymotion',
    "Dynamic Code Loading": r'loadClass|DexClassLoader',
    "Reflection Usage": r'reflection\.invoke|getMethod|getField',
    "JNI References": r'jni|JavaVM',
    "Third-Party SDKs": r'FacebookSDK|FirebaseSDK|GoogleAnalytics',
    "Telemetry References": r'telemetry|tracking|analytics',
    "Permission Requests": r'requestPermissions|grantUriPermission',
    "Broadcast Receivers": r'registerReceiver|BroadcastReceiver',
    "Activity Launching": r'startActivity|Intent\.FLAG_ACTIVITY',
    "Service Binding": r'bindService|ServiceConnection',
    "Alarm Manager Usage": r'AlarmManager\.set|setRepeating',
    "Notification Manager Usage": r'NotificationManager\.notify',
    "Content Resolver Usage": r'ContentResolver\.query',
    "Telephony Manager Usage": r'TelephonyManager\.getDeviceId',
    "Location Manager Usage": r'LocationManager\.requestLocationUpdates',
    "Camera Usage": r'Camera\.open|Camera\.startPreview',
    "SMS Sending": r'SmsManager\.sendTextMessage',
    "Call Management": r'TelephonyManager\.call',
    "Media Player Usage": r'MediaPlayer\.create|MediaPlayer\.start',
    "Clipboard Access": r'ClipboardManager\.setText|getText',
    "File I/O Operations": r'FileInputStream|FileOutputStream',
    "Network Operations": r'HttpURLConnection|OkHttpClient',
    "Encryption Libraries": r'AES|RSA|Cipher',
    "Hashing Algorithms": r'SHA1|SHA256|MD5',
    "Compression Libraries": r'GZIPOutputStream|Deflater',
    "HTTP Proxy Settings": r'Proxy\.setHttpProxy',
    "VPN Connections": r'VpnService\.prepare|Builder\.establish',
    "Biometric Authentication": r'FingerprintManager|BiometricPrompt',
    "Battery Optimization": r'PowerManager\.isIgnoringBatteryOptimizations',
    "Boot Receiver": r'BOOT_COMPLETED',
    "App Update Mechanism": r'updateApp|selfUpdate',
    "In-App Purchase Handling": r'InAppBillingService',
    "Push Notifications": r'FirebaseMessagingService|GCM',
    "Social Sharing": r'ShareCompat|ShareActionProvider',
    "Web Services Integration": r'Retrofit|Volley',
    "XML Parsing": r'SAXParser|XmlPullParser',
    "JSON Parsing": r'JSONObject|Gson|Jackson',
    "SQLite Operations": r'SQLiteOpenHelper|RoomDatabase',
    "File Encryption": r'FileCipher|EncryptedFile',
    "Remote Code Execution": r'eval|exec|Runtime\.exec',
    "Shell Commands": r'/system/bin/sh|Runtime\.exec',
    "Certificate Pinning": r'CertificatePinner|SSLSocketFactory',
    "Root Detection Methods": r'checkRootMethod|RootTools',
    "Screen Capture Prevention": r'FLAG_SECURE',
    "Key Store Usage": r'KeyStore\.getInstance|initKeyStore',
    "Crash Reporting": r'Crashlytics|FirebaseCrashlytics',
    "Analytics Integration": r'Mixpanel|Segment|Amplitude',
    "Ad Network Integration": r'AdMob|UnityAds|Vungle',
    "In-App Browser": r'WebView\.loadUrl|CustomTabsIntent',
    "Third-Party Libraries": r'OkHttp|Retrofit|Glide|Picasso',
    "License Verification": r'LicenseChecker|GooglePlayLicense',
    "Dynamic Feature Modules": r'DynamicDelivery|SplitInstallManager',
    "ProGuard/R8 Mappings": r'ProGuard|R8',
    "Certificate Validation": r'X509Certificate|TrustManager',
    "APK Signing": r'ApkSigner|signapk',
    "App Obfuscation": r'StringEncryption|ClassObfuscation',
    "Code Injection": r'Injection|Hooking',
    "Data Leakage": r'LeakCanary|StrictMode',
    "Application Sandboxing": r'Sandbox|Isolation',
    "Dependency Injection": r'Dagger|Koin',
    "Multithreading": r'Thread|AsyncTask|HandlerThread',
    "Reflection Obfuscation": r'ReflectionUsage|DynamicInvocation',
    "Custom Class Loaders": r'CustomClassLoader|DexClassLoader',
    "Resource Hacking": r'resource\.getIdentifier|AssetManager',
    "Anti-Tampering": r'TamperDetection|IntegrityCheck',
    "Root Commands": r'installPackage|su|busybox',
    "Dynamic Permissions": r'requestPermissions|onRequestPermissionsResult',
    "Foreground Services": r'ForegroundService|startForeground',
    "Gesture Detection": r'GestureDetector|MotionEvent',
    "Input Method Editors": r'InputMethodManager|Keyboard',
    "App Shortcuts": r'ShortcutManager|DynamicShortcuts',
    "Screen Orientation": r'SetRequestedOrientation|Orientation',
    "Live Wallpapers": r'WallpaperService|onCreateEngine',
    "Accessibility Services": r'AccessibilityService|AccessibilityEvent',
    "Voice Commands": r'VoiceInteractor|SpeechRecognizer',
    "Media Projection": r'MediaProjection|VirtualDisplay',
    "Job Scheduler": r'JobScheduler|JobInfo',
    "Work Manager": r'WorkManager|OneTimeWorkRequest',
    "Firebase Integration": r'FirebaseAuth|FirebaseDatabase',
    "Google Maps Integration": r'GoogleMap|MapView',
    "Bluetooth Operations": r'BluetoothAdapter|BluetoothDevice',
    "NFC Operations": r'NfcAdapter|NdefMessage',
    "USB Operations": r'UsbManager|UsbDevice',
    "Sensors Usage": r'SensorManager|SensorEvent',
    "Gesture Libraries": r'GestureLibrary|GestureOverlayView',
    "Push Notification Handling": r'FirebaseMessagingService|GcmListenerService',
    "App Widgets": r'AppWidgetProvider|AppWidgetManager',
    "Home Screen Launchers": r'LauncherActivity|MainActivity',
    "IMEI Retrieval": r'TelephonyManager\.getImei|getDeviceId',
    "MAC Address Retrieval": r'WifiInfo\.getMacAddress|getConnectionInfo',
    "SIM Card Information": r'TelephonyManager\.getSimCountryIso|getSimOperator',
    "Battery Information": r'BatteryManager\.getIntProperty|onBatteryChanged',
    "Power Management": r'PowerManager\.wakeUp|goToSleep',
    "Display Metrics": r'DisplayMetrics|getDisplayMetrics',
    "Clipboard Management": r'ClipboardManager\.setPrimaryClip|getPrimaryClip',
    "Accessibility Features": r'AccessibilityEvent|AccessibilityService',
    "Text-to-Speech": r'TextToSpeech|SpeechSynthesizer',
    "Speech Recognition": r'SpeechRecognizer|RecognizerIntent',
    "Notifications": r'NotificationCompat|NotificationChannel',
    "Camera Features": r'Camera2|CameraManager',
    "Media Recording": r'MediaRecorder|AudioRecord',
    "Video Playback": r'MediaPlayer|VideoView',
    "Audio Playback": r'AudioManager|MediaPlayer',
    "Bluetooth Features": r'BluetoothGatt|BluetoothProfile',
    "NFC Features": r'NfcAdapter|NdefRecord',
    "USB Features": r'UsbAccessory|UsbDeviceConnection',
    "Sensors Features": r'Gyroscope|Accelerometer',
    "Location Features": r'FusedLocationProviderClient|LocationRequest',
    "Gesture Features": r'GestureDetector|ScaleGestureDetector',
    "Vibration Features": r'Vibrator|VibrationEffect',
    "Network State": r'ConnectivityManager|NetworkInfo',
    "Wi-Fi Features": r'WifiManager|WifiConfiguration',
    "Mobile Data Features": r'TelephonyManager|DataEnabled',
    "GPS Features": r'GpsStatus|LocationListener',
    "Usernames": r'username\s*=\s*["\'][^"\']*["\']',
    "Passwords": r'password\s*=\s*["\'][^"\']*["\']',
    "API Keys": r'api[_-]?key\s*=\s*["\'][^"\']*["\']',
    "Token References": r'token\s*=\s*["\'][^"\']*["\']',
    "IP Addresses": r'\b(?:\d{1,3}\.){3}\d{1,3}\b',
    "Ports": r':(\d{1,5})\b',
    "Ports and IPs": r'\b(?:\d{1,3}\.){3}\d{1,3}:(\d{1,5})\b',
    # Advanced Detection Patterns
    "UEFI Injections": r'EFI/BOOT/BOOTX64.EFI|EFI/Microsoft/Boot/bootmgfw.efi',
    "Hardware Injections": r'I2C|SPI|UART|GPIO|DMA|PCIe',
    "Advanced Memory-Based Injections": r'MemoryMappedFile|Process Hollowing|Reflective DLL Injection|Fileless Malware Indicators',
    "Module-Based Injections": r'LoadLibrary|GetProcAddress|GetModuleHandle',
    "Advanced Tampering": r'SelfModification|CodeCave|NOP Sled|RunPE|Anti-Forensic Techniques',
    "Firmware Manipulation": r'BIOS|Firmware|ACPI|UEFI|PXE',
    "Code Signing Verification": r'CodeSignature|SignedCode',
    "Anti-VM Techniques": r'IsDebuggerPresent|CheckRemoteDebuggerPresent|VirtualMachine',
    "API Hooking Detection": r'Hooking|APIHook',
    "Sandbox Evasion Techniques": r'Sandbox|VMware|VirtualBox|Emulator',
    "Persistence Mechanisms": r'RunOnce|Startup|Registry Run|Scheduled Tasks',
    "Fileless Malware Indicators": r'PowerShell|WMI|Memory Injection',
    "Polymorphic Code Detection": r'Polymorphic|Metamorphic',
    "Metamorphic Code Analysis": r'Metamorphic|Mutation',
    "Rootkit Detection": r'Rootkit|HideProcess|HideModule',
    "Registry Monitoring (Windows)": r'Registry|HKLM|HKCU',
    "DLL Hijacking": r'DLLHijack|DllSearchOrder',
    "Command and Control (C2) Communications": r'C2|Command&Control|Beacon',
    "Privilege Escalation Indicators": r'PrivilegeEscalation|UAC bypass',
    "Data Exfiltration Patterns": r'DataExfiltration|DataLeak|FTP|HTTP|DNSExfil',
    "Heuristic-based Detection Patterns": r'Heuristic|Anomaly|SuspiciousBehavior',
    "Memory Scraping Indicators": r'MemoryScraping|CredentialDump',
    "Hardware ID Spoofing Patterns": r'HWIDSpoof|MACSpoof|IMIESpoof',
    "File Encryption Techniques": r'FileEncryption|AES|RSA|Cipher',
    "DLL Side-loading Patterns": r'DLLSideload|SideLoading',
    "Obfuscated Code Signatures": r'Obfuscated|StringEncryption|CodeObfuscation',
    "Polymorphic Shellcode Detection": r'PolymorphicShellcode|ShellcodeMutation',
    "Steganography Indicators": r'Steganography|ImageStego|AudioStego',
    "Cryptojacking Signatures": r'Cryptojacking|MiningScript|CoinMiner',
    "Exploit Kit Indicators": r'ExploitKit|EKScript|ExploitBundle',
    "Process Injection Variants": r'ProcessInjection|ThreadHijacking|CodeInjection',
    "MemoryTampering": r'MemoryTampering|MemoryCorruption',
    "Hidden Domains": r'\bhidden(?:\.\w+)+\b',
}

# ----------------------------- Modular Analysis Functions -----------------------------

def print_banner():
    banner = """
███████╗██╗      █████╗ ███╗   ███╗███████╗ █████╗ ██████╗ ██╗  ██╗██████╗ ██████╗ 
██╔════╝██║     ██╔══██╗████╗ ████║██╔════╝██╔══██╗██╔══██╗██║ ██╔╝██╔══██╗██╔══██╗
█████╗  ██║     ███████║██╔████╔██║█████╗  ███████║██████╔╝█████╔╝ ██████╔╝██████╔╝
██╔══╝  ██║     ██╔══██║██║╚██╔╝██║██╔══╝  ██╔══██║██╔═══╝ ██╔═██╗ ██╔═══╝ ██╔══██╗
██║     ███████╗██║  ██║██║ ╚═╝ ██║███████╗██║  ██║██║     ██║  ██╗██║     ██║  ██║
╚═╝     ╚══════╝╚═╝  ╚═╝╚═╝     ╚═╝╚══════╝╚═╝  ╚═╝╚═╝     ╚═╝  ╚═╝╚═╝     ╚═╝  ╚═╝
                                                                                 Pro
    A Comprehensive Multi-Platform Dynamic Analysis Tool
    Version: 1.0.0 | Author: Haroon Awan
    
Usage: FlameMaster_pro.py <application_file> [--custom "PatternName:Regex"]...
    """
    console.print(Panel(Text(banner, style="cyan"), style="bold blue"))

def unpack_file(file_path, output_dir):
    console.print(f"{INFO}[+] Unpacking In Progress{RESET}")
    logging.info(f"Unpacking {file_path} to {output_dir}")
    try:
        if zipfile.is_zipfile(file_path):
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                zip_ref.extractall(output_dir)
            console.print(f"{SUCCESS}[+] Unpacked to {output_dir}{RESET}")
            logging.info(f"Successfully unpacked to {output_dir}")
        else:
            console.print(f"{ERROR}[-] Unsupported archive format for {file_path}{RESET}")
            logging.error(f"Unsupported archive format for {file_path}")
            sys.exit(1)
    except zipfile.BadZipFile:
        console.print(f"{ERROR}[-] Failed to unzip {file_path}{RESET}")
        logging.error(f"Failed to unzip {file_path}")
        sys.exit(1)

def analyze_apk(apk_path, output_dir):
    console.print(f"{INFO}[>] Analyzing APK In Progress{RESET}")
    logging.info(f"Analyzing APK: {apk_path}")
    
    try:
        apk, _, dx = AnalyzeAPK(apk_path)
    except Exception as e:
        console.print(f"{ERROR}[-] Failed to analyze APK: {e}{RESET}")
        logging.error(f"Failed to analyze APK: {e}")
        sys.exit(1)
    
    # Create unzipped directory
    unzipped_dir = os.path.join(output_dir, "unzipped")
    os.makedirs(unzipped_dir, exist_ok=True)
    unpack_file(apk_path, unzipped_dir)
    
    # Save AndroidManifest.xml
    manifest_path = os.path.join(output_dir, "AndroidManifest.xml")
    try:
        manifest_data = apk.get_android_manifest_axml().get_xml()
        with open(manifest_path, 'wb') as f:
            f.write(manifest_data)
        
        console.print(f"{SUCCESS}[+] AndroidManifest.xml saved to {manifest_path}{RESET}")
        logging.info(f"Saved AndroidManifest.xml to {manifest_path}")
        
        # Verify well-formed XML
        try:
            tree = ET.parse(manifest_path)
            root = tree.getroot()
            console.print(f"{SUCCESS}[+] AndroidManifest.xml is well-formed.{RESET}")
            logging.info("AndroidManifest.xml is well-formed.")
        except ET.ParseError as pe:
            console.print(f"{ERROR}[-] Failed to parse AndroidManifest.xml: {pe}{RESET}")
            logging.error(f"Failed to parse AndroidManifest.xml: {pe}")
    except Exception as e:
        console.print(f"{WARNING}[!] Failed to save AndroidManifest.xml: {e}{RESET}")
        logging.warning(f"Failed to save AndroidManifest.xml: {e}")
    
    # Decompile DEX files
    dex_dir = os.path.join(output_dir, "classes.dex.out")
    os.makedirs(dex_dir, exist_ok=True)
    console.print(f"{INFO}[>] Decompiling DEX files{RESET}")
    logging.info(f"Decompiling DEX files to {dex_dir}")
    
    try:
        # Get all methods
        methods = list(dx.get_methods())
        total_methods = len(methods)
        console.print(f"{INFO}[>] Total methods to decompile: {total_methods}{RESET}")
        logging.info(f"Total methods to decompile: {total_methods}")

        with Progress(SpinnerColumn(), BarColumn(), TextColumn("{task.description}")) as progress:
            task = progress.add_task("Decompiling methods...", total=total_methods)
            for method in methods:
                try:
                    class_name = method.get_class_name().replace('/', '_').replace('$', '_')
                    method_name = method.get_name().replace('<', '').replace('>', '')
                    filename = f"{class_name}_{method_name}.txt"
                    filepath = os.path.join(dex_dir, filename)
                    
                    # Get source code
                    source_code = method.get_source()
                    if source_code:
                        with open(filepath, 'w', encoding='utf-8') as f:
                            f.write(source_code)
                    else:
                        # Try to get disassembly
                        disasm = method.get_disas()
                        if disasm:
                            with open(filepath, 'w', encoding='utf-8') as f:
                                f.write(str(disasm))
                        else:
                            logging.warning(f"No source or disassembly for method: {method}")
                except AttributeError as ae:
                    logging.warning(f"Attribute error for method {method}: {ae}")
                except Exception as e:
                    logging.warning(f"Failed to write {filepath}: {e}")
                progress.update(task, advance=1)
        
        console.print(f"{SUCCESS}[+] Decompiled code saved to {dex_dir}{RESET}")
        logging.info(f"Decompiled code saved to {dex_dir}")
    except Exception as e:
        console.print(f"{ERROR}[-] Failed to decompile DEX files: {e}{RESET}")
        logging.error(f"Failed to decompile DEX files: {e}")
    
    return apk, dx

def analyze_ipa(ipa_path, output_dir):
    console.print(f"{INFO}[>] Analyzing IPA In Progress{RESET}")
    logging.info(f"Analyzing IPA: {ipa_path}")
    
    # Create temporary extraction directory
    extract_dir = os.path.join(output_dir, "extracted")
    os.makedirs(extract_dir, exist_ok=True)
    
    try:
        with zipfile.ZipFile(ipa_path, 'r') as zip_ref:
            zip_ref.extractall(extract_dir)
        console.print(f"{SUCCESS}[+] IPA unpacked to {extract_dir}{RESET}")
        logging.info(f"Successfully unpacked IPA to {extract_dir}")
    except zipfile.BadZipFile:
        console.print(f"{ERROR}[-] Failed to unzip {ipa_path}{RESET}")
        logging.error(f"Failed to unzip {ipa_path}")
        sys.exit(1)
    
    # Locate Info.plist and binary
    plist_path = None
    binary_path = None
    
    for root, dirs, files in os.walk(extract_dir):
        if 'Info.plist' in files:
            plist_path = os.path.join(root, 'Info.plist')
        for file in files:
            if not file.endswith(('.png', '.jpg', '.json', '.plist', '.strings')):
                potential_binary = os.path.join(root, file)
                if os.access(potential_binary, os.X_OK):
                    binary_path = potential_binary
                    break
    
    if plist_path:
        console.print(f"{SUCCESS}[+] Info.plist found at {plist_path}{RESET}")
        logging.info(f"Info.plist found at {plist_path}")
        
        # Parse Info.plist using plistlib
        try:
            import plistlib
            with open(plist_path, 'rb') as f:
                plist_data = plistlib.load(f)
            
            plist_output = os.path.join(output_dir, "Info_parsed.txt")
            with open(plist_output, 'w', encoding='utf-8') as f:
                for key, value in plist_data.items():
                    f.write(f"{key}: {value}\n")
            
            console.print(f"{SUCCESS}[+] Parsed Info.plist saved to {plist_output}{RESET}")
            logging.info(f"Parsed Info.plist saved to {plist_output}")
        except Exception as e:
            console.print(f"{WARNING}[!] Failed to parse Info.plist: {e}{RESET}")
            logging.warning(f"Failed to parse Info.plist: {e}")
    else:
        console.print(f"{WARNING}[!] Info.plist not found{RESET}")
        logging.warning("Info.plist not found")
    
    # Analyze binary if found
    if binary_path:
        console.print(f"{SUCCESS}[+] Binary found at {binary_path}{RESET}")
        logging.info(f"Binary found at {binary_path}")
        
        try:
            binary = lief.parse(binary_path)
            if binary:
                binary_info = os.path.join(output_dir, "binary_analysis.txt")
                with open(binary_info, 'w', encoding='utf-8') as f:
                    f.write(f"Binary Format: {binary.format}\n")
                    f.write(f"Architecture: {binary.header.cpu_type}\n")
                    f.write(f"Number of Sections: {len(binary.sections)}\n")
                    f.write(f"Number of Segments: {len(binary.segments)}\n")
                    
                    if hasattr(binary, 'imported_functions'):
                        f.write("\nImported Functions:\n")
                        for imp in binary.imported_functions:
                            f.write(f"  {imp}\n")
                
                console.print(f"{SUCCESS}[+] Binary analysis saved to {binary_info}{RESET}")
                logging.info(f"Binary analysis saved to {binary_info}")
        except Exception as e:
            console.print(f"{WARNING}[!] Failed to analyze binary with LIEF: {e}{RESET}")
            logging.warning(f"Failed to analyze binary with LIEF: {e}")
        
        # Generate hex dump
        hex_output = os.path.join(output_dir, "binary_hex_dump.txt")
        try:
            with open(binary_path, 'rb') as bin_file, open(hex_output, 'w') as hex_file:
                offset = 0
                while True:
                    chunk = bin_file.read(16)
                    if not chunk:
                        break
                    hex_str = ' '.join(f'{b:02x}' for b in chunk)
                    ascii_str = ''.join(chr(b) if 32 <= b < 127 else '.' for b in chunk)
                    hex_file.write(f'{offset:08x}: {hex_str:<48} {ascii_str}\n')
                    offset += 16
            
            console.print(f"{SUCCESS}[+] Hex dump saved to {hex_output}{RESET}")
            logging.info(f"Hex dump saved to {hex_output}")
        except Exception as e:
            console.print(f"{WARNING}[!] Failed to create hex dump: {e}{RESET}")
            logging.warning(f"Failed to create hex dump: {e}")
    else:
        console.print(f"{WARNING}[!] Binary not found{RESET}")
        logging.warning("Binary not found")

def analyze_executable(file_path, output_dir, os_type='Windows'):
    console.print(f"{INFO}[>] Analyzing Executable In Progress{RESET}")
    logging.info(f"Analyzing Executable: {file_path} as {os_type}")
    
    # Static Analysis using LIEF
    try:
        binary = lief.parse(file_path)
        if not binary:
            raise Exception("Failed to parse the binary.")
        
        # Extract Headers Information
        headers_output = os.path.join(output_dir, f"{os_type}_headers.txt")
        with open(headers_output, 'w', encoding='utf-8') as f:
            f.write(f"=== {os_type} Executable Headers ===\n")
            
            if binary.format == lief.EXE_FORMATS.PE:
                f.write(f"PE Machine: {binary.header.machine}\n")
                f.write(f"Characteristics: {hex(binary.header.characteristics)}\n")
                f.write(f"Number of Sections: {len(binary.sections)}\n")
                f.write(f"Number of Symbols: {len(binary.symbols)}\n")
                
                for section in binary.sections:
                    f.write(f"\nSection: {section.name}\n")
                    f.write(f"  Size: {section.size} bytes\n")
                    f.write(f"  Virtual Size: {section.virtual_size} bytes\n")
                    f.write(f"  Characteristics: {hex(section.characteristics)}\n")
                    
            elif binary.format == lief.EXE_FORMATS.ELF:
                f.write(f"ELF Class: {binary.header.identity_class}\n")
                f.write(f"ELF Type: {binary.header.file_type}\n")
                f.write(f"Machine Type: {binary.header.machine_type}\n")
                f.write(f"Number of Sections: {len(binary.sections)}\n")
                
                for section in binary.sections:
                    f.write(f"\nSection: {section.name}\n")
                    f.write(f"  Type: {section.type}\n")
                    f.write(f"  Size: {section.size} bytes\n")
                    f.write(f"  Virtual Address: {hex(section.virtual_address)}\n")
                    
            elif binary.format == lief.EXE_FORMATS.MACHO:
                f.write(f"MACH-O CPU: {binary.header.cpu_type}\n")
                f.write(f"MACH-O Subtype: {binary.header.cpu_subtype}\n")
                f.write(f"Number of Commands: {binary.header.nb_cmds}\n")
                f.write(f"Number of Sections: {len(binary.sections)}\n")
                
                for section in binary.sections:
                    f.write(f"\nSection: {section.name}\n")
                    f.write(f"  Size: {section.size} bytes\n")
                    f.write(f"  Offset: {section.offset}\n")
                    
            else:
                f.write("Unsupported executable format.\n")
        
        console.print(f"{SUCCESS}[+] Headers extracted to {headers_output}{RESET}")
        logging.info(f"Headers extracted to {headers_output}")
        
        # Extract Imports/Exports
        imports_output = os.path.join(output_dir, f"{os_type}_imports.txt")
        with open(imports_output, 'w', encoding='utf-8') as f:
            f.write(f"=== {os_type} Executable Imports ===\n")
            
            if binary.format == lief.EXE_FORMATS.PE:
                if hasattr(binary, 'imports'):
                    for imp in binary.imports:
                        f.write(f"\nImported DLL: {imp.name}\n")
                        if hasattr(imp, 'entries'):
                            for func in imp.entries:
                                if hasattr(func, 'name') and func.name:
                                    f.write(f"  Function: {func.name}\n")
                                elif hasattr(func, 'iat_value'):
                                    f.write(f"  Function (by ordinal): {hex(func.iat_value)}\n")
            elif binary.format == lief.EXE_FORMATS.ELF:
                if hasattr(binary, 'dynamic_symbols'):
                    for sym in binary.dynamic_symbols:
                        if sym.imported:
                            f.write(f"Imported Symbol: {sym.name}\n")
            elif binary.format == lief.EXE_FORMATS.MACHO:
                if hasattr(binary, 'imported_functions'):
                    for func in binary.imported_functions:
                        f.write(f"Imported Function: {func}\n")
        
        console.print(f"{SUCCESS}[+] Imports extracted to {imports_output}{RESET}")
        logging.info(f"Imports extracted to {imports_output}")
        
        # Advanced Static Analysis
        advanced_static_analysis(binary, output_dir)
        
    except Exception as e:
        console.print(f"{ERROR}[-] Static analysis failed: {e}{RESET}")
        logging.error(f"Static analysis failed: {e}")
    
    # Dynamic Analysis (Optional - can be enabled with flag)
    if os_type.lower() in ['windows', 'linux']:
        console.print(f"{INFO}[>] Starting Dynamic Analysis for {os_type}{RESET}")
        logging.info(f"Starting Dynamic Analysis for {os_type}")
        
        # Extract strings from binary
        strings_output = os.path.join(output_dir, "extracted_strings.txt")
        try:
            strings_cmd = ["strings", file_path]
            result = subprocess.run(strings_cmd, capture_output=True, text=True)
            if result.returncode == 0:
                with open(strings_output, 'w', encoding='utf-8') as f:
                    f.write(result.stdout)
                console.print(f"{SUCCESS}[+] Strings extracted to {strings_output}{RESET}")
                logging.info(f"Strings extracted to {strings_output}")
            else:
                console.print(f"{WARNING}[!] 'strings' command failed, using Python fallback{RESET}")
                # Fallback to Python implementation
                with open(file_path, 'rb') as f:
                    data = f.read()
                    strings = re.findall(rb'[ -~]{4,}', data)
                    with open(strings_output, 'w', encoding='utf-8') as f_out:
                        for s in strings:
                            f_out.write(s.decode('utf-8', errors='ignore') + '\n')
                console.print(f"{SUCCESS}[+] Strings extracted (fallback) to {strings_output}{RESET}")
        except Exception as e:
            console.print(f"{WARNING}[!] Failed to extract strings: {e}{RESET}")
            logging.warning(f"Failed to extract strings: {e}")
    else:
        console.print(f"{WARNING}[!] Dynamic analysis not fully supported for {os_type}{RESET}")
        logging.warning(f"Dynamic analysis not fully supported for {os_type}")

def analyze_golang_binary(file_path, output_dir):
    console.print(f"{INFO}[>] Analyzing Golang Binary In Progress{RESET}")
    logging.info(f"Analyzing Golang Binary: {file_path}")
    
    try:
        # Extract string literals
        with open(file_path, 'rb') as f:
            content = f.read()
            strings = re.findall(rb'[ -~]{4,}', content)
            strings = [s.decode('utf-8', errors='ignore') for s in strings]
        
        strings_output = os.path.join(output_dir, "golang_strings.txt")
        with open(strings_output, 'w', encoding='utf-8') as f:
            f.write('\n'.join(strings))
        
        console.print(f"{SUCCESS}[+] Extracted strings saved to {strings_output}{RESET}")
        logging.info(f"Extracted strings saved to {strings_output}")
        
        # Look for Go-specific patterns
        go_patterns = {
            "Go Runtime Functions": r'runtime\.',
            "Go Packages": r'package\s+\w+',
            "Go Imports": r'import\s+\([^)]+\)',
            "Go Functions": r'func\s+\w+\s*\([^)]*\)',
        }
        
        go_analysis = os.path.join(output_dir, "go_specific_patterns.txt")
        with open(go_analysis, 'w', encoding='utf-8') as f:
            for pattern_name, pattern in go_patterns.items():
                matches = re.findall(pattern, content.decode('utf-8', errors='ignore'))
                if matches:
                    f.write(f"{pattern_name}:\n")
                    for match in matches[:10]:  # Limit to first 10 matches
                        f.write(f"  {match}\n")
                    f.write("\n")
        
        console.print(f"{SUCCESS}[+] Go-specific patterns saved to {go_analysis}{RESET}")
        logging.info(f"Go-specific patterns saved to {go_analysis}")
        
    except Exception as e:
        console.print(f"{WARNING}[!] Failed to analyze Golang binary: {e}{RESET}")
        logging.warning(f"Failed to analyze Golang binary: {e}")

def analyze_pcap(pcap_path, output_dir):
    console.print(f"{INFO}[>] Analyzing PCAP In Progress{RESET}")
    logging.info(f"Analyzing PCAP: {pcap_path}")
    
    try:
        packets = rdpcap(pcap_path)
        packet_count = len(packets)
        
        pcap_output = os.path.join(output_dir, "pcap_summary.txt")
        with open(pcap_output, 'w', encoding='utf-8') as f:
            f.write(f"Total Packets: {packet_count}\n")
            f.write(f"File Size: {os.path.getsize(pcap_path)} bytes\n")
            f.write(f"Duration: {packets[-1].time - packets[0].time:.2f} seconds\n\n")
            
            # Count protocols
            protocols = {}
            for pkt in packets:
                if hasattr(pkt, 'proto'):
                    proto = pkt.proto
                    protocols[proto] = protocols.get(proto, 0) + 1
            
            f.write("Protocol Distribution:\n")
            for proto, count in protocols.items():
                f.write(f"  {proto}: {count} packets\n")
            
            f.write("\nPacket Summary (first 50 packets):\n")
            for i, pkt in enumerate(packets[:50]):
                f.write(f"{i+1}: {pkt.summary()}\n")
        
        console.print(f"{SUCCESS}[+] PCAP summary saved to {pcap_output}{RESET}")
        logging.info(f"PCAP summary saved to {pcap_output}")
        
        # Extract HTTP requests if any
        http_output = os.path.join(output_dir, "http_requests.txt")
        try:
            with open(http_output, 'w', encoding='utf-8') as f:
                for pkt in packets:
                    if pkt.haslayer('TCP') and pkt.haslayer('Raw'):
                        try:
                            payload = pkt['Raw'].load.decode('utf-8', errors='ignore')
                            if 'HTTP' in payload:
                                f.write(payload[:500] + "\n" + "="*50 + "\n")
                        except:
                            continue
            console.print(f"{SUCCESS}[+] HTTP requests saved to {http_output}{RESET}")
        except Exception as e:
            console.print(f"{WARNING}[!] Failed to extract HTTP requests: {e}{RESET}")
            
    except Exception as e:
        console.print(f"{ERROR}[-] Failed to analyze PCAP: {e}{RESET}")
        logging.error(f"Failed to analyze PCAP: {e}")

def analyze_document(file_path, output_dir):
    console.print(f"{INFO}[>] Analyzing Document In Progress{RESET}")
    logging.info(f"Analyzing Document: {file_path}")
    
    try:
        if file_path.endswith('.pdf'):
            text = extract_text(file_path)
            text_output = os.path.join(output_dir, "document_text.txt")
            with open(text_output, 'w', encoding='utf-8') as f:
                f.write(text)
            console.print(f"{SUCCESS}[+] Extracted text saved to {text_output}{RESET}")
            logging.info(f"Extracted text saved to {text_output}")
            
            # Extract metadata
            try:
                from pdfminer.pdfparser import PDFParser
                from pdfminer.pdfdocument import PDFDocument
                with open(file_path, 'rb') as fp:
                    parser = PDFParser(fp)
                    doc = PDFDocument(parser)
                    metadata = doc.info[0] if doc.info else {}
                    
                    meta_output = os.path.join(output_dir, "pdf_metadata.txt")
                    with open(meta_output, 'w', encoding='utf-8') as f:
                        for key, value in metadata.items():
                            f.write(f"{key}: {value}\n")
                    console.print(f"{SUCCESS}[+] PDF metadata saved to {meta_output}{RESET}")
            except Exception as e:
                console.print(f"{WARNING}[!] Failed to extract PDF metadata: {e}{RESET}")
                
        elif file_path.endswith('.docx'):
            doc = Document(file_path)
            text = '\n'.join([para.text for para in doc.paragraphs])
            text_output = os.path.join(output_dir, "document_text.txt")
            with open(text_output, 'w', encoding='utf-8') as f:
                f.write(text)
            console.print(f"{SUCCESS}[+] Extracted text saved to {text_output}{RESET}")
            logging.info(f"Extracted text saved to {text_output}")
            
            # Extract core properties
            core_props = doc.core_properties
            meta_output = os.path.join(output_dir, "docx_metadata.txt")
            with open(meta_output, 'w', encoding='utf-8') as f:
                f.write(f"Author: {core_props.author}\n")
                f.write(f"Created: {core_props.created}\n")
                f.write(f"Modified: {core_props.modified}\n")
                f.write(f"Title: {core_props.title}\n")
                f.write(f"Subject: {core_props.subject}\n")
            console.print(f"{SUCCESS}[+] DOCX metadata saved to {meta_output}{RESET}")
        else:
            console.print(f"{WARNING}[!] Unsupported document format: {file_path}{RESET}")
            logging.warning(f"Unsupported document format: {file_path}")
    except Exception as e:
        console.print(f"{ERROR}[-] Failed to analyze document: {e}{RESET}")
        logging.error(f"Failed to analyze document: {e}")

def analyze_archive(file_path, output_dir):
    console.print(f"{INFO}[>] Analyzing Archive In Progress{RESET}")
    logging.info(f"Analyzing Archive: {file_path}")
    
    # Create extraction directory
    extract_dir = os.path.join(output_dir, "extracted")
    os.makedirs(extract_dir, exist_ok=True)
    
    try:
        if file_path.endswith('.zip'):
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                zip_ref.extractall(extract_dir)
                files = zip_ref.namelist()
                
        elif file_path.endswith('.rar'):
            try:
                import rarfile
                with rarfile.RarFile(file_path, 'r') as rar_ref:
                    rar_ref.extractall(extract_dir)
                    files = rar_ref.namelist()
            except ImportError:
                console.print(f"{ERROR}[-] rarfile module not installed. Install with: pip install rarfile{RESET}")
                return
                
        elif file_path.endswith('.7z'):
            try:
                import py7zr
                with py7zr.SevenZipFile(file_path, 'r') as sz_ref:
                    sz_ref.extractall(extract_dir)
                    files = sz_ref.getnames()
            except ImportError:
                console.print(f"{ERROR}[-] py7zr module not installed. Install with: pip install py7zr{RESET}")
                return
        
        # List extracted files
        archive_info = os.path.join(output_dir, "archive_contents.txt")
        with open(archive_info, 'w', encoding='utf-8') as f:
            f.write(f"Archive: {file_path}\n")
            f.write(f"Extracted to: {extract_dir}\n")
            f.write("\nContents:\n")
            for file in files:
                f.write(f"  {file}\n")
        
        console.print(f"{SUCCESS}[+] Archive extracted to {extract_dir}{RESET}")
        console.print(f"{SUCCESS}[+] Archive contents saved to {archive_info}{RESET}")
        logging.info(f"Archive extracted to {extract_dir}")
        
        # Analyze extracted files
        for root, dirs, files in os.walk(extract_dir):
            for file in files:
                file_path_full = os.path.join(root, file)
                file_ext = os.path.splitext(file)[1].lower()
                
                # Skip if file is too large (> 10MB)
                if os.path.getsize(file_path_full) > 10 * 1024 * 1024:
                    continue
                
                # Analyze based on file type
                if file_ext in ['.exe', '.dll', '.bin']:
                    analyze_executable(file_path_full, output_dir, 'Windows')
                elif file_ext in ['.elf']:
                    analyze_executable(file_path_full, output_dir, 'Linux')
                elif file_ext in ['.pdf', '.docx']:
                    analyze_document(file_path_full, output_dir)
        
    except Exception as e:
        console.print(f"{ERROR}[-] Failed to analyze archive: {e}{RESET}")
        logging.error(f"Failed to analyze archive: {e}")

def analyze_powershell(file_path, output_dir):
    console.print(f"{INFO}[>] Analyzing PowerShell Script In Progress{RESET}")
    logging.info(f"Analyzing PowerShell Script: {file_path}")
    
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        # Define suspicious patterns specific to PowerShell
        suspicious_patterns = {
            "Invoke-Expression": r'Invoke-Expression',
            "DownloadString": r'DownloadString',
            "Add-Type": r'Add-Type',
            "IEX": r'\bIEX\b',
            "EncodedCommand": r'-EncodedCommand\s+\S+',
            "Invoke-WebRequest": r'Invoke-WebRequest',
            "Start-Process": r'Start-Process',
            "New-Object": r'New-Object',
            "Net.WebClient": r'Net\.WebClient',
            "Bypass Execution Policy": r'-ExecutionPolicy\s+Bypass',
            "MemoryInjection": r'MemoryMappedFile|Process Hollowing|Reflective DLL Injection',
            "UEFIManipulation": r'EFI/BOOT/BOOTX64.EFI|EFI/Microsoft/Boot/bootmgfw.efi',
        }
        
        results = {}
        for pattern_name, pattern in suspicious_patterns.items():
            matches = re.findall(pattern, content, re.IGNORECASE)
            if matches:
                results[pattern_name] = matches
        
        # Display results
        if results:
            table = Table(title="Suspicious Patterns in PowerShell Script", show_lines=True)
            table.add_column("Pattern", style="cyan", no_wrap=True)
            table.add_column("Matches", style="magenta")
            for pattern, matches in results.items():
                table.add_row(pattern, ', '.join(set(matches[:5])))  # Show first 5 unique matches
            console.print(table)
            
            # Save detailed results
            ps_output = os.path.join(output_dir, "powershell_analysis.txt")
            with open(ps_output, 'w', encoding='utf-8') as f:
                f.write("PowerShell Script Analysis\n")
                f.write("=" * 50 + "\n\n")
                for pattern, matches in results.items():
                    f.write(f"{pattern}:\n")
                    for match in set(matches):
                        f.write(f"  - {match}\n")
                    f.write("\n")
            
            console.print(f"{SUCCESS}[+] PowerShell analysis saved to {ps_output}{RESET}")
            logging.info(f"Suspicious patterns found: {results}")
        else:
            console.print(f"{SUCCESS}[+] No suspicious patterns found in PowerShell script.{RESET}")
            logging.info("No suspicious patterns found in PowerShell script.")
        
        # Extract URLs and IPs
        url_pattern = r'https?://[^\s<>"\']+|www\.[^\s<>"\']+'
        ip_pattern = r'\b(?:\d{1,3}\.){3}\d{1,3}\b'
        
        urls = re.findall(url_pattern, content)
        ips = re.findall(ip_pattern, content)
        
        if urls or ips:
            network_output = os.path.join(output_dir, "powershell_network_info.txt")
            with open(network_output, 'w', encoding='utf-8') as f:
                if urls:
                    f.write("URLs Found:\n")
                    for url in set(urls):
                        f.write(f"  {url}\n")
                    f.write("\n")
                if ips:
                    f.write("IP Addresses Found:\n")
                    for ip in set(ips):
                        f.write(f"  {ip}\n")
            
            console.print(f"{SUCCESS}[+] Network information saved to {network_output}{RESET}")
            
    except Exception as e:
        console.print(f"{ERROR}[-] Failed to analyze PowerShell script: {e}{RESET}")
        logging.error(f"Failed to analyze PowerShell script: {e}")

def analyze_email(file_path, output_dir):
    console.print(f"{INFO}[>] Analyzing E-Mail File In Progress{RESET}")
    logging.info(f"Analyzing E-Mail File: {file_path}")
    
    try:
        import email
        from email import policy
        from email.parser import BytesParser
        
        with open(file_path, 'rb') as f:
            msg = BytesParser(policy=policy.default).parse(f)
        
        subject = msg.get('Subject', 'N/A')
        from_ = msg.get('From', 'N/A')
        to = msg.get('To', 'N/A')
        date = msg.get('Date', 'N/A')
        
        email_output = os.path.join(output_dir, "email_details.txt")
        with open(email_output, 'w', encoding='utf-8') as f:
            f.write("Email Analysis Report\n")
            f.write("=" * 50 + "\n\n")
            f.write(f"Subject: {subject}\n")
            f.write(f"From: {from_}\n")
            f.write(f"To: {to}\n")
            f.write(f"Date: {date}\n\n")
            
            # Headers
            f.write("Headers:\n")
            for key, value in msg.items():
                f.write(f"  {key}: {value}\n")
            
            # Body
            f.write("\nBody:\n")
            if msg.is_multipart():
                for part in msg.iter_parts():
                    if part.get_content_type() == "text/plain":
                        body = part.get_content()
                        f.write(body)
                        break
            else:
                body = msg.get_content()
                f.write(str(body))
        
        console.print(f"{SUCCESS}[+] E-Mail details saved to {email_output}{RESET}")
        logging.info(f"E-Mail details saved to {email_output}")
        
        # Extract attachments if any
        if msg.is_multipart():
            attachment_dir = os.path.join(output_dir, "attachments")
            os.makedirs(attachment_dir, exist_ok=True)
            
            attachment_count = 0
            for part in msg.iter_parts():
                if part.get_filename():
                    filename = part.get_filename()
                    attachment_path = os.path.join(attachment_dir, filename)
                    with open(attachment_path, 'wb') as f:
                        f.write(part.get_content())
                    attachment_count += 1
            
            if attachment_count > 0:
                console.print(f"{SUCCESS}[+] {attachment_count} attachments saved to {attachment_dir}{RESET}")
        
    except Exception as e:
        console.print(f"{ERROR}[-] Failed to analyze E-Mail file: {e}{RESET}")
        logging.error(f"Failed to analyze E-Mail file: {e}")

def analyze_uefi_injections(file_path, output_dir):
    console.print(f"{INFO}[>] Analyzing UEFI Injections In Progress{RESET}")
    logging.info(f"Analyzing UEFI Injections in: {file_path}")
    
    try:
        with open(file_path, 'rb') as f:
            content = f.read()
            text_content = content.decode('utf-8', errors='ignore')
            matches = re.findall(PATTERNS["UEFI Injections"], text_content)
            
            if matches:
                uefi_output = os.path.join(output_dir, "uefi_injections.txt")
                with open(uefi_output, 'w', encoding='utf-8') as f_out:
                    f_out.write("UEFI Injection Patterns Found:\n")
                    f_out.write("=" * 50 + "\n\n")
                    for match in set(matches):
                        f_out.write(f"{match}\n")
                
                console.print(f"{SUCCESS}[+] UEFI Injections detected and saved to {uefi_output}{RESET}")
                logging.info(f"UEFI Injections detected and saved to {uefi_output}")
            else:
                console.print(f"{SUCCESS}[+] No UEFI Injections detected.{RESET}")
                logging.info("No UEFI Injections detected.")
    except Exception as e:
        console.print(f"{WARNING}[!] Failed to analyze UEFI injections: {e}{RESET}")
        logging.warning(f"Failed to analyze UEFI injections: {e}")

def analyze_hardware_injections(file_path, output_dir):
    console.print(f"{INFO}[>] Analyzing Hardware Injections In Progress{RESET}")
    logging.info(f"Analyzing Hardware Injections in: {file_path}")
    
    try:
        with open(file_path, 'rb') as f:
            content = f.read()
            text_content = content.decode('utf-8', errors='ignore')
            matches = re.findall(PATTERNS["Hardware Injections"], text_content)
            
            if matches:
                hardware_output = os.path.join(output_dir, "hardware_injections.txt")
                with open(hardware_output, 'w', encoding='utf-8') as f_out:
                    f_out.write("Hardware Injection Patterns Found:\n")
                    f_out.write("=" * 50 + "\n\n")
                    for match in set(matches):
                        f_out.write(f"{match}\n")
                
                console.print(f"{SUCCESS}[+] Hardware Injections detected and saved to {hardware_output}{RESET}")
                logging.info(f"Hardware Injections detected and saved to {hardware_output}")
            else:
                console.print(f"{SUCCESS}[+] No Hardware Injections detected.{RESET}")
                logging.info("No Hardware Injections detected.")
    except Exception as e:
        console.print(f"{WARNING}[!] Failed to analyze Hardware injections: {e}{RESET}")
        logging.warning(f"Failed to analyze Hardware injections: {e}")

def analyze_memory_injections(file_path, output_dir):
    console.print(f"{INFO}[>] Analyzing Advanced Memory-Based Injections In Progress{RESET}")
    logging.info(f"Analyzing Advanced Memory-Based Injections in: {file_path}")
    
    try:
        with open(file_path, 'rb') as f:
            content = f.read()
            text_content = content.decode('utf-8', errors='ignore')
            pattern = PATTERNS["Advanced Memory-Based Injections"]
            matches = re.findall(pattern, text_content)
            
            if matches:
                memory_output = os.path.join(output_dir, "memory_injections.txt")
                with open(memory_output, 'w', encoding='utf-8') as f_out:
                    f_out.write("Advanced Memory-Based Injection Patterns Found:\n")
                    f_out.write("=" * 50 + "\n\n")
                    for match in set(matches):
                        f_out.write(f"{match}\n")
                
                console.print(f"{SUCCESS}[+] Advanced Memory-Based Injections detected and saved to {memory_output}{RESET}")
                logging.info(f"Advanced Memory-Based Injections detected and saved to {memory_output}")
            else:
                console.print(f"{SUCCESS}[+] No Advanced Memory-Based Injections detected.{RESET}")
                logging.info("No Advanced Memory-Based Injections detected.")
    except Exception as e:
        console.print(f"{WARNING}[!] Failed to analyze Advanced Memory-Based injections: {e}{RESET}")
        logging.warning(f"Failed to analyze Advanced Memory-Based injections: {e}")

def analyze_module_injections(file_path, output_dir):
    console.print(f"{INFO}[>] Analyzing Module-Based Injections In Progress{RESET}")
    logging.info(f"Analyzing Module-Based Injections in: {file_path}")
    
    try:
        with open(file_path, 'rb') as f:
            content = f.read()
            text_content = content.decode('utf-8', errors='ignore')
            matches = re.findall(PATTERNS["Module-Based Injections"], text_content)
            
            if matches:
                module_output = os.path.join(output_dir, "module_injections.txt")
                with open(module_output, 'w', encoding='utf-8') as f_out:
                    f_out.write("Module-Based Injection Patterns Found:\n")
                    f_out.write("=" * 50 + "\n\n")
                    for match in set(matches):
                        f_out.write(f"{match}\n")
                
                console.print(f"{SUCCESS}[+] Module-Based Injections detected and saved to {module_output}{RESET}")
                logging.info(f"Module-Based Injections detected and saved to {module_output}")
            else:
                console.print(f"{SUCCESS}[+] No Module-Based Injections detected.{RESET}")
                logging.info("No Module-Based Injections detected.")
    except Exception as e:
        console.print(f"{WARNING}[!] Failed to analyze Module-Based injections: {e}{RESET}")
        logging.warning(f"Failed to analyze Module-Based injections: {e}")

def analyze_advanced_tampering(file_path, output_dir):
    console.print(f"{INFO}[>] Analyzing Advanced Tampering In Progress{RESET}")
    logging.info(f"Analyzing Advanced Tampering in: {file_path}")
    
    try:
        with open(file_path, 'rb') as f:
            content = f.read()
            text_content = content.decode('utf-8', errors='ignore')
            matches = re.findall(PATTERNS["Advanced Tampering"], text_content)
            
            if matches:
                tampering_output = os.path.join(output_dir, "advanced_tampering.txt")
                with open(tampering_output, 'w', encoding='utf-8') as f_out:
                    f_out.write("Advanced Tampering Patterns Found:\n")
                    f_out.write("=" * 50 + "\n\n")
                    for match in set(matches):
                        f_out.write(f"{match}\n")
                
                console.print(f"{SUCCESS}[+] Advanced Tampering detected and saved to {tampering_output}{RESET}")
                logging.info(f"Advanced Tampering detected and saved to {tampering_output}")
            else:
                console.print(f"{SUCCESS}[+] No Advanced Tampering detected.{RESET}")
                logging.info("No Advanced Tampering detected.")
    except Exception as e:
        console.print(f"{WARNING}[!] Failed to analyze Advanced Tampering: {e}{RESET}")
        logging.warning(f"Failed to analyze Advanced Tampering: {e}")

def analyze_firmware_manipulation(file_path, output_dir):
    console.print(f"{INFO}[>] Analyzing Firmware Manipulation In Progress{RESET}")
    logging.info(f"Analyzing Firmware Manipulation in: {file_path}")
    
    try:
        with open(file_path, 'rb') as f:
            content = f.read()
            text_content = content.decode('utf-8', errors='ignore')
            matches = re.findall(PATTERNS["Firmware Manipulation"], text_content)
            
            if matches:
                firmware_output = os.path.join(output_dir, "firmware_manipulation.txt")
                with open(firmware_output, 'w', encoding='utf-8') as f_out:
                    f_out.write("Firmware Manipulation Patterns Found:\n")
                    f_out.write("=" * 50 + "\n\n")
                    for match in set(matches):
                        f_out.write(f"{match}\n")
                
                console.print(f"{SUCCESS}[+] Firmware Manipulation detected and saved to {firmware_output}{RESET}")
                logging.info(f"Firmware Manipulation detected and saved to {firmware_output}")
            else:
                console.print(f"{SUCCESS}[+] No Firmware Manipulation detected.{RESET}")
                logging.info("No Firmware Manipulation detected.")
    except Exception as e:
        console.print(f"{WARNING}[!] Failed to analyze Firmware Manipulation: {e}{RESET}")
        logging.warning(f"Failed to analyze Firmware Manipulation: {e}")

def analyze_process_injections(file_path, output_dir):
    console.print(f"{INFO}[>] Analyzing Process Injections In Progress{RESET}")
    logging.info(f"Analyzing Process Injections in: {file_path}")
    
    try:
        with open(file_path, 'rb') as f:
            content = f.read()
            text_content = content.decode('utf-8', errors='ignore')
            matches = re.findall(PATTERNS["Process Injection Variants"], text_content)
            
            if matches:
                process_output = os.path.join(output_dir, "process_injections.txt")
                with open(process_output, 'w', encoding='utf-8') as f_out:
                    f_out.write("Process Injection Patterns Found:\n")
                    f_out.write("=" * 50 + "\n\n")
                    for match in set(matches):
                        f_out.write(f"{match}\n")
                
                console.print(f"{SUCCESS}[+] Process Injections detected and saved to {process_output}{RESET}")
                logging.info(f"Process Injections detected and saved to {process_output}")
            else:
                console.print(f"{SUCCESS}[+] No Process Injections detected.{RESET}")
                logging.info("No Process Injections detected.")
    except Exception as e:
        console.print(f"{WARNING}[!] Failed to analyze Process Injections: {e}{RESET}")
        logging.warning(f"Failed to analyze Process Injections: {e}")

def analyze_firmware(file_path, output_dir):
    console.print(f"{INFO}[>] Analyzing Firmware Injections In Progress{RESET}")
    logging.info(f"Analyzing Firmware Injections in: {file_path}")
    
    try:
        with open(file_path, 'rb') as f:
            content = f.read()
            text_content = content.decode('utf-8', errors='ignore')
            matches = re.findall(PATTERNS["Firmware Manipulation"], text_content)
            
            if matches:
                firmware_output = os.path.join(output_dir, "firmware_injections.txt")
                with open(firmware_output, 'w', encoding='utf-8') as f_out:
                    f_out.write("Firmware Injection Patterns Found:\n")
                    f_out.write("=" * 50 + "\n\n")
                    for match in set(matches):
                        f_out.write(f"{match}\n")
                
                console.print(f"{SUCCESS}[+] Firmware Injections detected and saved to {firmware_output}{RESET}")
                logging.info(f"Firmware Injections detected and saved to {firmware_output}")
            else:
                console.print(f"{SUCCESS}[+] No Firmware Injections detected.{RESET}")
                logging.info("No Firmware Injections detected.")
    except Exception as e:
        console.print(f"{WARNING}[!] Failed to analyze Firmware Injections: {e}{RESET}")
        logging.warning(f"Failed to analyze Firmware Injections: {e}")

def advanced_static_analysis(binary, output_dir):
    """
    Perform advanced static analysis on the binary to detect sophisticated malware behaviors.
    """
    console.print(f"{INFO}[>] Performing Advanced Static Analysis{RESET}")
    logging.info("Starting advanced static analysis")
    
    try:
        # Detect Obfuscation Techniques
        obfuscation_patterns = [
            r'packer',
            r'encrypt',
            r'decrypt',
            r'obfuscate',
            r'protected',
            r'cipher',
            r'crypt',
            r'polymorphic',
            r'metamorphic',
        ]
        
        obfuscation_output = os.path.join(output_dir, "advanced_static_analysis.txt")
        with open(obfuscation_output, 'w', encoding='utf-8') as f:
            f.write("Advanced Static Analysis Report\n")
            f.write("=" * 50 + "\n\n")
            
            # Check section names for obfuscation
            f.write("Obfuscation Analysis:\n")
            obfuscation_found = False
            for section in binary.sections:
                section_name = section.name.lower()
                for pattern in obfuscation_patterns:
                    if re.search(pattern, section_name):
                        f.write(f"  Suspect section: {section.name}\n")
                        obfuscation_found = True
                        break
            
            if not obfuscation_found:
                f.write("  No obvious obfuscation detected in section names.\n")
            
            # Check for anti-debugging imports
            f.write("\nAnti-Debugging Analysis:\n")
            anti_debug_patterns = [
                'IsDebuggerPresent',
                'CheckRemoteDebuggerPresent',
                'OutputDebugString',
                'NtQueryInformationProcess',
            ]
            
            anti_debug_found = False
            if hasattr(binary, 'imports'):
                for imp in binary.imports:
                    for pattern in anti_debug_patterns:
                        if hasattr(imp, 'name') and imp.name and pattern.lower() in imp.name.lower():
                            f.write(f"  Anti-debug import: {imp.name}\n")
                            anti_debug_found = True
            
            if not anti_debug_found:
                f.write("  No anti-debugging imports detected.\n")
            
            # Check for suspicious APIs
            f.write("\nSuspicious API Analysis:\n")
            suspicious_apis = [
                'VirtualAlloc',
                'VirtualProtect',
                'WriteProcessMemory',
                'CreateRemoteThread',
                'LoadLibrary',
                'GetProcAddress',
            ]
            
            suspicious_found = False
            if hasattr(binary, 'imports'):
                for imp in binary.imports:
                    if hasattr(imp, 'entries'):
                        for entry in imp.entries:
                            if hasattr(entry, 'name') and entry.name:
                                for api in suspicious_apis:
                                    if api.lower() in entry.name.lower():
                                        f.write(f"  Suspicious API: {entry.name}\n")
                                        suspicious_found = True
            
            if not suspicious_found:
                f.write("  No obviously suspicious APIs detected.\n")
            
            # Check for packing indicators
            f.write("\nPacking Indicators:\n")
            packing_indicators = [
                ('Small number of imports', len(binary.imports) < 10 if hasattr(binary, 'imports') else False),
                ('High entropy sections', False),  # Would need entropy calculation
                ('Unusual section names', any('UPX' in s.name for s in binary.sections) if hasattr(binary, 'sections') else False),
            ]
            
            for indicator, condition in packing_indicators:
                if condition:
                    f.write(f"  {indicator}: YES\n")
                else:
                    f.write(f"  {indicator}: NO\n")
            
            # Entry point analysis
            f.write("\nEntry Point Analysis:\n")
            if hasattr(binary, 'entrypoint'):
                f.write(f"  Entry Point: {hex(binary.entrypoint)}\n")
            else:
                f.write("  Entry Point: Not available\n")
        
        console.print(f"{SUCCESS}[+] Advanced static analysis saved to {obfuscation_output}{RESET}")
        logging.info("Advanced static analysis completed successfully")
        
    except Exception as e:
        console.print(f"{WARNING}[!] Advanced static analysis failed: {e}{RESET}")
        logging.warning(f"Advanced static analysis failed: {e}")

def extract_endpoints_and_urls(directory, output_dir):
    console.print(f"{INFO}[>] Extracting Endpoints and URLs{RESET}")
    logging.info(f"Extracting endpoints and URLs from {directory}")
    
    patterns = {
        "Endpoint References": PATTERNS["Endpoint References"],
        "URLs": PATTERNS["URL References"],
        "HTTP References": PATTERNS["HTTP References"],
        "API References": PATTERNS["API References"],
    }
    
    results = search_patterns(directory, patterns)
    endpoints_output = os.path.join(output_dir, "endpoints_urls.txt")
    
    try:
        with open(endpoints_output, 'w', encoding='utf-8') as f:
            f.write("Endpoints and URLs Found:\n")
            f.write("=" * 50 + "\n\n")
            
            for pattern, matches in results.items():
                if matches:
                    f.write(f"{pattern}:\n")
                    unique_matches = set()
                    for file_path, match in matches:
                        # Clean up the match
                        match_clean = match.strip('"\'').strip()
                        unique_matches.add(match_clean)
                    
                    for match in sorted(unique_matches):
                        f.write(f"  {match}\n")
                    f.write("\n")
        
        console.print(f"{SUCCESS}[+] Endpoints and URLs extracted to {endpoints_output}{RESET}")
        logging.info(f"Endpoints and URLs extracted to {endpoints_output}")
    except Exception as e:
        console.print(f"{WARNING}[!] Failed to write endpoints and URLs: {e}{RESET}")
        logging.warning(f"Failed to write endpoints and URLs: {e}")

def extract_sql_information(directory, output_dir):
    console.print(f"{INFO}[>] Extracting SQL Information{RESET}")
    logging.info(f"Extracting SQL information from {directory}")
    
    patterns = {
        "SQL Database References": PATTERNS["SQL Database References"],
        "SQLite References": PATTERNS["SQLite References"],
    }
    
    results = search_patterns(directory, patterns)
    sql_output = os.path.join(output_dir, "sql_information.txt")
    
    try:
        with open(sql_output, 'w', encoding='utf-8') as f:
            f.write("SQL Information Found:\n")
            f.write("=" * 50 + "\n\n")
            
            for pattern, matches in results.items():
                if matches:
                    f.write(f"{pattern}:\n")
                    for file_path, match in matches:
                        f.write(f"  File: {file_path}\n")
                        f.write(f"  Match: {match}\n\n")
        
        console.print(f"{SUCCESS}[+] SQL Information extracted to {sql_output}{RESET}")
        logging.info(f"SQL Information extracted to {sql_output}")
    except Exception as e:
        console.print(f"{WARNING}[!] Failed to write SQL information: {e}{RESET}")
        logging.warning(f"Failed to write SQL information: {e}")

def extract_hidden_domains(directory, output_dir):
    console.print(f"{INFO}[>] Extracting Hidden Domains{RESET}")
    logging.info(f"Extracting hidden domains from {directory}")
    
    patterns = {
        "Hidden Domains": PATTERNS["Hidden Domains"],
        "Subdomain References": PATTERNS["Subdomain References"],
        "Domain References": PATTERNS["Domain References"],
    }
    
    results = search_patterns(directory, patterns)
    hidden_domains_output = os.path.join(output_dir, "hidden_domains.txt")
    
    try:
        with open(hidden_domains_output, 'w', encoding='utf-8') as f:
            f.write("Hidden Domains Found:\n")
            f.write("=" * 50 + "\n\n")
            
            for pattern, matches in results.items():
                if matches:
                    f.write(f"{pattern}:\n")
                    unique_matches = set()
                    for file_path, match in matches:
                        unique_matches.add(match.strip('"\'').strip())
                    
                    for match in sorted(unique_matches):
                        f.write(f"  {match}\n")
                    f.write("\n")
        
        console.print(f"{SUCCESS}[+] Hidden Domains extracted to {hidden_domains_output}{RESET}")
        logging.info(f"Hidden Domains extracted to {hidden_domains_output}")
    except Exception as e:
        console.print(f"{WARNING}[!] Failed to write hidden domains: {e}{RESET}")
        logging.warning(f"Failed to write hidden domains: {e}")

def extract_user_credentials(directory, output_dir):
    console.print(f"{INFO}[>] Extracting User Credentials{RESET}")
    logging.info(f"Extracting user credentials from {directory}")
    
    patterns = {
        "Hardcoded Secrets": PATTERNS["Hardcoded Secrets"],
        "Usernames": PATTERNS["Usernames"],
        "Passwords": PATTERNS["Passwords"],
    }
    
    results = search_patterns(directory, patterns)
    credentials_output = os.path.join(output_dir, "user_credentials.txt")
    
    try:
        with open(credentials_output, 'w', encoding='utf-8') as f:
            f.write("User Credentials Found:\n")
            f.write("=" * 50 + "\n\n")
            f.write("WARNING: This file may contain sensitive information!\n\n")
            
            for pattern, matches in results.items():
                if matches:
                    f.write(f"{pattern}:\n")
                    for file_path, match in matches:
                        # Truncate match if too long
                        match_display = match[:100] + "..." if len(match) > 100 else match
                        f.write(f"  File: {file_path}\n")
                        f.write(f"  Match: {match_display}\n\n")
        
        console.print(f"{SUCCESS}[+] User credentials extracted to {credentials_output}{RESET}")
        logging.info(f"User credentials extracted to {credentials_output}")
    except Exception as e:
        console.print(f"{WARNING}[!] Failed to write user credentials: {e}{RESET}")
        logging.warning(f"Failed to write user credentials: {e}")

def extract_ip_ports(directory, output_dir):
    console.print(f"{INFO}[>] Extracting IP Addresses and Ports{RESET}")
    logging.info(f"Extracting IP addresses and ports from {directory}")
    
    patterns = {
        "Valid IP Addresses": PATTERNS["Valid IP Addresses"],
        "Ports and IPs": PATTERNS["Ports and IPs"],
        "Socket References": PATTERNS["Socket References"],
    }
    
    results = search_patterns(directory, patterns)
    ip_ports_output = os.path.join(output_dir, "ip_ports.txt")
    
    try:
        with open(ip_ports_output, 'w', encoding='utf-8') as f:
            f.write("IP Addresses and Ports Found:\n")
            f.write("=" * 50 + "\n\n")
            
            for pattern, matches in results.items():
                if matches:
                    f.write(f"{pattern}:\n")
                    unique_matches = set()
                    for file_path, match in matches:
                        unique_matches.add(match)
                    
                    for match in sorted(unique_matches):
                        f.write(f"  {match}\n")
                    f.write("\n")
        
        console.print(f"{SUCCESS}[+] IP Addresses and Ports extracted to {ip_ports_output}{RESET}")
        logging.info(f"IP Addresses and Ports extracted to {ip_ports_output}")
    except Exception as e:
        console.print(f"{WARNING}[!] Failed to write IP Addresses and Ports: {e}{RESET}")
        logging.warning(f"Failed to write IP Addresses and Ports: {e}")

def extract_api_information(directory, output_dir):
    console.print(f"{INFO}[>] Extracting API Information{RESET}")
    logging.info(f"Extracting API information from {directory}")
    
    patterns = {
        "API References": PATTERNS["API References"],
        "API Keys": PATTERNS["API Keys"],
        "Token References": PATTERNS["Token References"],
        "API Variable References": PATTERNS["API Variable References"],
    }
    
    results = search_patterns(directory, patterns)
    api_output = os.path.join(output_dir, "api_information.txt")
    
    try:
        with open(api_output, 'w', encoding='utf-8') as f:
            f.write("API Information Found:\n")
            f.write("=" * 50 + "\n\n")
            f.write("WARNING: This file may contain sensitive API keys!\n\n")
            
            for pattern, matches in results.items():
                if matches:
                    f.write(f"{pattern}:\n")
                    for file_path, match in matches:
                        # Truncate sensitive data
                        if "key" in pattern.lower() or "token" in pattern.lower():
                            match_display = match[:50] + "..." if len(match) > 50 else match
                        else:
                            match_display = match
                        f.write(f"  File: {file_path}\n")
                        f.write(f"  Match: {match_display}\n\n")
        
        console.print(f"{SUCCESS}[+] API Information extracted to {api_output}{RESET}")
        logging.info(f"API Information extracted to {api_output}")
    except Exception as e:
        console.print(f"{WARNING}[!] Failed to write API Information: {e}{RESET}")
        logging.warning(f"Failed to write API Information: {e}")

def extract_folder_paths(directory, output_dir):
    console.print(f"{INFO}[>] Extracting Folder Paths{RESET}")
    logging.info(f"Extracting folder paths from {directory}")
    
    patterns = {
        "Folder Paths": r'(?:[A-Za-z]:\\|/)(?:[^\\/:*?"<>|\r\n]+[\\/])*[^\\/:*?"<>|\r\n]*',
        "File References": PATTERNS["File References"],
    }
    
    results = search_patterns(directory, patterns)
    folders_output = os.path.join(output_dir, "folder_paths.txt")
    
    try:
        with open(folders_output, 'w', encoding='utf-8') as f:
            f.write("Folder Paths and File References Found:\n")
            f.write("=" * 50 + "\n\n")
            
            for pattern, matches in results.items():
                if matches:
                    f.write(f"{pattern}:\n")
                    unique_matches = set()
                    for file_path, match in matches:
                        unique_matches.add(match)
                    
                    for match in sorted(unique_matches):
                        f.write(f"  {match}\n")
                    f.write("\n")
        
        console.print(f"{SUCCESS}[+] Folder Paths extracted to {folders_output}{RESET}")
        logging.info(f"Folder Paths extracted to {folders_output}")
    except Exception as e:
        console.print(f"{WARNING}[!] Failed to write Folder Paths: {e}{RESET}")
        logging.warning(f"Failed to write Folder Paths: {e}")

# ----------------------------- Utility Functions -----------------------------

def search_patterns(directory, patterns, file_extension=None):
    console.print(f"{INFO}[>] Searching for patterns in {directory}{RESET}")
    logging.info(f"Searching for patterns in {directory}")
    
    results = {pattern: [] for pattern in patterns.keys()}
    compiled_patterns = {name: re.compile(pattern, re.IGNORECASE) for name, pattern in patterns.items()}
    
    if not os.path.exists(directory):
        logging.warning(f"Directory does not exist: {directory}")
        return results
    
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file_extension and not file.endswith(file_extension):
                continue
            
            file_path = os.path.join(root, file)
            try:
                # Skip binary files
                if is_binary_file(file_path):
                    continue
                
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                    for name, pattern in compiled_patterns.items():
                        for match in pattern.findall(content):
                            # Handle tuple matches
                            if isinstance(match, tuple):
                                match = ' '.join(match)
                            if match:  # Only add non-empty matches
                                results[name].append((file_path, match))
            except Exception as e:
                logging.debug(f"Failed to read {file_path}: {e}")
                continue
    
    return results

def is_binary_file(file_path):
    """Check if a file is binary."""
    try:
        with open(file_path, 'rb') as f:
            chunk = f.read(1024)
            return b'\x00' in chunk
    except:
        return False

def display_results(results):
    if not any(results.values()):
        console.print(f"{INFO}[>] No patterns found in the search results{RESET}")
        return
    
    table = Table(title="Pattern Search Results", show_lines=True)
    table.add_column("Pattern", style="cyan", no_wrap=True)
    table.add_column("File", style="magenta")
    table.add_column("Match", style="green")
    
    row_count = 0
    for pattern, matches in results.items():
        if matches:
            for file_path, match in matches[:5]:  # Show only first 5 matches per pattern
                display_match = (match[:75] + '...') if len(match) > 75 else match
                table.add_row(pattern, os.path.basename(file_path), str(display_match))
                row_count += 1
                if row_count >= 50:  # Limit total rows
                    break
        if row_count >= 50:
            break
    
    if row_count > 0:
        console.print(table)
        console.print(f"{INFO}[>] Showing {row_count} matches (truncated){RESET}")
    else:
        console.print(f"{INFO}[>] No matches found{RESET}")

def generate_html_report(results, output_dir, custom_results=None):
    env = Environment(loader=FileSystemLoader('.'))
    template = env.from_string("""
    <!DOCTYPE html>
    <html>
    <head>
        <title>FlameMaster Pro Analysis Report</title>
        <style>
            body { 
                font-family: Arial, sans-serif; 
                background-color: #f4f4f4;
                margin: 20px;
                line-height: 1.6;
            }
            .container {
                max-width: 1200px;
                margin: 0 auto;
                background: white;
                padding: 20px;
                border-radius: 5px;
                box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            }
            h1 {
                color: #333;
                border-bottom: 2px solid #4CAF50;
                padding-bottom: 10px;
            }
            h2 {
                color: #444;
                margin-top: 30px;
                border-bottom: 1px solid #ddd;
                padding-bottom: 5px;
            }
            table {
                border-collapse: collapse;
                width: 100%;
                margin: 20px 0;
                font-size: 14px;
            }
            th {
                background-color: #4CAF50;
                color: white;
                padding: 12px;
                text-align: left;
            }
            td {
                border: 1px solid #ddd;
                padding: 8px;
                vertical-align: top;
            }
            tr:nth-child(even) {
                background-color: #f9f9f9;
            }
            tr:hover {
                background-color: #f5f5f5;
            }
            .pattern-name {
                font-weight: bold;
                color: #0066cc;
            }
            .filename {
                font-family: monospace;
                color: #666;
            }
            .match {
                font-family: monospace;
                word-break: break-all;
            }
            .timestamp {
                color: #888;
                font-size: 12px;
                text-align: right;
                margin-top: 20px;
            }
            .warning {
                background-color: #fff3cd;
                border: 1px solid #ffeaa7;
                padding: 15px;
                margin: 20px 0;
                border-radius: 4px;
            }
        </style>
    </head>
    <body>
        <div class="container">
            <h1>🔥 FlameMaster Pro Analysis Report</h1>
            <div class="timestamp">Generated on: {{ timestamp }}</div>
            
            {% if any_results %}
            <div class="warning">
                <strong>⚠️ Security Notice:</strong> This report may contain sensitive information including API keys, passwords, and IP addresses. Handle with care.
            </div>
            
            <h2>Pattern Search Results</h2>
            <table>
                <tr>
                    <th>Pattern</th>
                    <th>File</th>
                    <th>Match</th>
                </tr>
                {% for pattern, matches in results.items() %}
                    {% if matches %}
                        {% for file, match in matches %}
                        <tr>
                            <td class="pattern-name">{{ pattern }}</td>
                            <td class="filename">{{ file }}</td>
                            <td class="match">{{ match }}</td>
                        </tr>
                        {% endfor %}
                    {% endif %}
                {% endfor %}
            </table>
            {% endif %}
            
            {% if custom_results and any_custom_results %}
            <h2>Custom Information Extraction</h2>
            <table>
                <tr>
                    <th>Custom Pattern</th>
                    <th>File</th>
                    <th>Match</th>
                </tr>
                {% for pattern, matches in custom_results.items() %}
                    {% if matches %}
                        {% for file, match in matches %}
                        <tr>
                            <td class="pattern-name">{{ pattern }}</td>
                            <td class="filename">{{ file }}</td>
                            <td class="match">{{ match }}</td>
                        </tr>
                        {% endfor %}
                    {% endif %}
                {% endfor %}
            </table>
            {% endif %}
            
            {% if not any_results and not any_custom_results %}
            <p>No patterns were found in the analysis.</p>
            {% endif %}
        </div>
    </body>
    </html>
    """)
    
    # Prepare data for template
    any_results = any(bool(matches) for matches in results.values())
    any_custom_results = custom_results and any(bool(matches) for matches in custom_results.values())
    
    html_content = template.render(
        results=results,
        custom_results=custom_results,
        any_results=any_results,
        any_custom_results=any_custom_results,
        timestamp=datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    )
    
    report_path = os.path.join(output_dir, "FlameMaster_pro_comprehensive_report.html")
    try:
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        console.print(f"{SUCCESS}[+] Comprehensive HTML report generated at {report_path}{RESET}")
        logging.info(f"Comprehensive HTML report generated at {report_path}")
    except Exception as e:
        console.print(f"{WARNING}[!] Failed to generate HTML report: {e}{RESET}")
        logging.warning(f"Failed to generate HTML report: {e}")

def display_manifest_details(apk, output_dir):
    manifest_path = os.path.join(output_dir, "AndroidManifest.xml")
    if os.path.exists(manifest_path):
        console.print(f"{INFO}[>] Displaying AndroidManifest.xml Details{RESET}")
        try:
            tree = ET.parse(manifest_path)
            root = tree.getroot()
            
            # Extract namespace
            ns = {'android': 'http://schemas.android.com/apk/res/android'}
            
            package = root.attrib.get('package', 'N/A')
            console.print(f"{SUCCESS}[+] Package: {package}{RESET}")
            logging.info(f"Package: {package}")
            
            # Find elements
            activities = root.findall(".//activity")
            services = root.findall(".//service")
            providers = root.findall(".//provider")
            receivers = root.findall(".//receiver")
            permissions = root.findall(".//uses-permission")
            
            # Display counts
            console.print(f"{INFO}[+] Activities: {len(activities)}{RESET}")
            console.print(f"{INFO}[+] Services: {len(services)}{RESET}")
            console.print(f"{INFO}[+] Providers: {len(providers)}{RESET}")
            console.print(f"{INFO}[+] Receivers: {len(receivers)}{RESET}")
            console.print(f"{INFO}[+] Permissions: {len(permissions)}{RESET}")
            
            # Save detailed manifest analysis
            manifest_analysis = os.path.join(output_dir, "manifest_analysis.txt")
            with open(manifest_analysis, 'w', encoding='utf-8') as f:
                f.write("Android Manifest Analysis\n")
                f.write("=" * 50 + "\n\n")
                f.write(f"Package: {package}\n\n")
                
                f.write("Permissions:\n")
                for perm in permissions:
                    name = perm.attrib.get(f'{{{ns["android"]}}}name', 'N/A')
                    f.write(f"  {name}\n")
                
                f.write("\nActivities:\n")
                for activity in activities:
                    name = activity.attrib.get(f'{{{ns["android"]}}}name', 'N/A')
                    exported = activity.attrib.get(f'{{{ns["android"]}}}exported', 'N/A')
                    f.write(f"  {name} (exported: {exported})\n")
                
                f.write("\nServices:\n")
                for service in services:
                    name = service.attrib.get(f'{{{ns["android"]}}}name', 'N/A')
                    exported = service.attrib.get(f'{{{ns["android"]}}}exported', 'N/A')
                    f.write(f"  {name} (exported: {exported})\n")
            
            console.print(f"{SUCCESS}[+] Manifest analysis saved to {manifest_analysis}{RESET}")
            logging.info("Displayed AndroidManifest.xml details")
            
        except Exception as e:
            console.print(f"{WARNING}[!] Failed to parse AndroidManifest.xml: {e}{RESET}")
            logging.warning(f"Failed to parse AndroidManifest.xml: {e}")
    else:
        console.print(f"{WARNING}[!] AndroidManifest.xml not found{RESET}")
        logging.warning("AndroidManifest.xml not found")

def extract_strings(apk_path, output_dir):
    strings_output = os.path.join(output_dir, "strings.txt")
    console.print(f"{INFO}[>] Extracting strings to {strings_output}{RESET}")
    logging.info(f"Extracting strings from {apk_path} to {strings_output}")
    
    try:
        # Try to use androguard first
        try:
            apk = APK(apk_path)
            strings = apk.get_strings()
            if strings:
                with open(strings_output, 'w', encoding='utf-8') as f:
                    f.write('\n'.join(strings))
                console.print(f"{SUCCESS}[+] Strings extracted to {strings_output}{RESET}")
                logging.info(f"Strings extracted to {strings_output}")
                return
        except:
            pass
        
        # Fallback: extract strings from all files
        with open(strings_output, 'w', encoding='utf-8') as f:
            # Unzip and extract strings
            with zipfile.ZipFile(apk_path, 'r') as zip_ref:
                for file_info in zip_ref.infolist():
                    if file_info.file_size < 10 * 1024 * 1024:  # Skip files larger than 10MB
                        try:
                            content = zip_ref.read(file_info.filename)
                            strings = re.findall(rb'[ -~]{4,}', content)
                            if strings:
                                f.write(f"\n=== {file_info.filename} ===\n")
                                for s in strings[:100]:  # Limit to first 100 strings per file
                                    f.write(s.decode('utf-8', errors='ignore') + '\n')
                        except:
                            continue
        
        console.print(f"{SUCCESS}[+] Strings extracted (fallback) to {strings_output}{RESET}")
        logging.info(f"Strings extracted to {strings_output}")
        
    except Exception as e:
        console.print(f"{WARNING}[!] Failed to extract strings: {e}{RESET}")
        logging.warning(f"Failed to extract strings: {e}")

def hex_dump(file_path, output_dir):
    hex_output = os.path.join(output_dir, f"{os.path.basename(file_path)}_hex_dump.txt")
    console.print(f"{INFO}[>] Creating Hex Dump at {hex_output}{RESET}")
    logging.info(f"Creating hex dump for {file_path} at {hex_output}")
    
    try:
        with open(file_path, 'rb') as bin_file, open(hex_output, 'w') as hex_file:
            offset = 0
            chunk_size = 16
            
            while True:
                chunk = bin_file.read(chunk_size)
                if not chunk:
                    break
                
                # Hex representation
                hex_str = ' '.join(f'{b:02x}' for b in chunk)
                hex_str = hex_str.ljust(chunk_size * 3 - 1)
                
                # ASCII representation
                ascii_str = ''.join(chr(b) if 32 <= b < 127 else '.' for b in chunk)
                
                hex_file.write(f'{offset:08x}: {hex_str}  {ascii_str}\n')
                offset += chunk_size
        
        console.print(f"{SUCCESS}[+] Hex dump created at {hex_output}{RESET}")
        logging.info(f"Hex dump created at {hex_output}")
    except Exception as e:
        console.print(f"{WARNING}[!] Failed to create hex dump: {e}{RESET}")
        logging.warning(f"Failed to create hex dump: {e}")

def custom_info_extractor(directory, custom_patterns, file_extension=None):
    console.print(f"{INFO}[>] Running Custom Information Extraction{RESET}")
    logging.info("Running custom information extraction")
    
    results = {}
    for pattern_name in custom_patterns.keys():
        results[pattern_name] = []
    
    compiled_custom_patterns = {name: re.compile(pattern, re.IGNORECASE) for name, pattern in custom_patterns.items()}
    
    if not os.path.exists(directory):
        logging.warning(f"Directory does not exist: {directory}")
        return results
    
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file_extension and not file.endswith(file_extension):
                continue
            
            file_path = os.path.join(root, file)
            try:
                if is_binary_file(file_path):
                    continue
                
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                    for name, pattern in compiled_custom_patterns.items():
                        for match in pattern.findall(content):
                            if isinstance(match, tuple):
                                match = ' '.join(match)
                            if match:
                                results[name].append((file_path, match))
            except Exception as e:
                logging.debug(f"Failed to read {file_path}: {e}")
                continue
    
    return results

def integrate_debugger(binary_path):
    console.print(f"{INFO}[>] Integrating with Debugger{RESET}")
    logging.info("Integrating with debugger")
    
    try:
        file_ext = os.path.splitext(binary_path)[1].lower()
        
        if sys.platform == "win32":
            debugger = 'ollydbg'  # Or x64dbg
            console.print(f"{INFO}[>] Suggested debugger for Windows: {debugger}{RESET}")
            console.print(f"{INFO}[>] Please open {binary_path} in {debugger} manually{RESET}")
        elif sys.platform == "linux":
            debugger = 'gdb'
            console.print(f"{INFO}[>] Starting GDB for {binary_path}{RESET}")
            try:
                subprocess.Popen(['gdb', binary_path])
                console.print(f"{SUCCESS}[+] Launched {debugger} for {binary_path}{RESET}")
            except FileNotFoundError:
                console.print(f"{WARNING}[!] GDB not found. Please install it with: sudo apt-get install gdb{RESET}")
        elif sys.platform == "darwin":
            debugger = 'lldb'
            console.print(f"{INFO}[>] Starting LLDB for {binary_path}{RESET}")
            try:
                subprocess.Popen(['lldb', binary_path])
                console.print(f"{SUCCESS}[+] Launched {debugger} for {binary_path}{RESET}")
            except FileNotFoundError:
                console.print(f"{WARNING}[!] LLDB not found. Please install Xcode command line tools{RESET}")
        
        logging.info(f"Debugger integration attempted for {binary_path}")
    except Exception as e:
        console.print(f"{WARNING}[!] Failed to integrate with debugger: {e}{RESET}")
        logging.warning(f"Failed to integrate with debugger: {e}")

def generate_comprehensive_report(results, custom_results, output_dir):
    generate_html_report(results, output_dir, custom_results)

def remediation_suggestions(results):
    suggestions = []
    
    if not any(results.values()):
        suggestions.append("No security issues detected in the pattern search.")
        return suggestions
    
    for pattern, matches in results.items():
        if pattern in ["Hardcoded Secrets", "Usernames", "Passwords", "API Keys", "Token References"] and matches:
            suggestions.append(f"Remove hardcoded secrets found in {pattern}. Consider using secure storage solutions like Android Keystore or iOS Keychain.")
        
        elif pattern in ["SSL Certificate Pinning"] and matches:
            suggestions.append("Implement proper SSL certificate pinning to prevent man-in-the-middle attacks.")
        
        elif pattern in ["Root Detection"] and matches:
            suggestions.append("Ensure your app gracefully handles rooted devices. Consider providing limited functionality rather than blocking entirely.")
        
        elif pattern in ["Emulator Detection"] and matches:
            suggestions.append("Avoid relying solely on emulator detection as it can be bypassed by sophisticated attackers.")
        
        elif pattern in ["Command Execution References", "Shell Commands"] and matches:
            suggestions.append("Avoid executing shell commands from the application. If necessary, validate and sanitize all inputs.")
        
        elif pattern in ["SQL Database References"] and matches:
            suggestions.append("Use parameterized queries to prevent SQL injection attacks.")
        
        elif pattern in ["File References"] and matches:
            suggestions.append("Validate all file paths and restrict file access to appropriate directories.")
        
        elif pattern in ["URL References", "Endpoint References"] and matches:
            suggestions.append("Use HTTPS for all network communications and validate server certificates.")
        
        elif pattern in ["WebView References", "JavaScript Interface"] and matches:
            suggestions.append("Disable JavaScript in WebView if not needed, or carefully sanitize all JavaScript interactions.")
        
        elif pattern in ["Advanced Memory-Based Injections", "Process Injection Variants"] and matches:
            suggestions.append("Implement stringent memory protection and monitor critical functions to prevent memory tampering.")
        
        elif pattern in ["UEFI Injections", "Firmware Manipulation"] and matches:
            suggestions.append("Secure firmware components and implement integrity checks to prevent firmware-level injections.")
    
    if not suggestions:
        suggestions.append("No specific remediation suggestions based on the detected patterns.")
    
    return suggestions

# ----------------------------- Main Functionality -----------------------------

def main():
    parser = ArgumentParser(
        description="FlameMaster Pro: A Modular Multi-Platform Analysis Tool for Cybersecurity",
        formatter_class=RawTextHelpFormatter,
        epilog="""
Examples:
  %(prog)s sample.apk
  %(prog)s malware.exe --custom "C2Server:http://[^\\s]+" --custom "Mutex:[A-Za-z0-9_]+"
  %(prog)s suspicious.pdf
  %(prog)s traffic.pcap
        """
    )
    parser.add_argument("file", help="Path to the application file to analyze")
    parser.add_argument("--custom", action='append', help="Custom pattern in the format 'PatternName:Regex'")
    parser.add_argument("--debug", action='store_true', help="Enable debug mode with more verbose output")
    parser.add_argument("--no-report", action='store_true', help="Skip HTML report generation")
    
    args = parser.parse_args()
    
    file_path = args.file
    if not os.path.isfile(file_path):
        console.print(f"{ERROR}[-] File {file_path} does not exist.{RESET}")
        logging.error(f"File {file_path} does not exist.")
        sys.exit(1)
    
    # Set debug logging if requested
    if args.debug:
        logging.getLogger().setLevel(logging.DEBUG)
        console.print(f"{INFO}[>] Debug mode enabled{RESET}")
    
    # Handle custom patterns if provided
    custom_patterns = {}
    if args.custom:
        for pattern_input in args.custom:
            if ':' in pattern_input:
                name, pattern = pattern_input.split(':', 1)
                custom_patterns[name.strip()] = pattern.strip()
                console.print(f"{INFO}[+] Added custom pattern: {name.strip()}{RESET}")
            else:
                console.print(f"{WARNING}[!] Invalid custom pattern format: {pattern_input}{RESET}")
                logging.warning(f"Invalid custom pattern format: {pattern_input}")
    
    file_ext = os.path.splitext(file_path)[1].lower()
    base_name = os.path.basename(file_path)
    base_dir = os.path.join(os.getcwd(), f"{base_name}_analysis_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}")
    os.makedirs(base_dir, exist_ok=True)
    
    print_banner()
    console.print(f"{INFO}[>] Starting analysis of: {file_path}{RESET}")
    console.print(f"{INFO}[>] Output directory: {base_dir}{RESET}")
    logging.info(f"Starting analysis of: {file_path}")
    logging.info(f"Output directory: {base_dir}")
    
    # Dictionary to map file extensions to their analysis functions and arguments
    analysis_map = {
        '.apk': {'function': analyze_apk, 'args': {}},
        '.ipa': {'function': analyze_ipa, 'args': {}},
        '.exe': {'function': analyze_executable, 'args': {'os_type': 'Windows'}},
        '.dll': {'function': analyze_executable, 'args': {'os_type': 'Windows'}},
        '.msi': {'function': analyze_executable, 'args': {'os_type': 'Windows'}},
        '.bin': {'function': analyze_executable, 'args': {'os_type': 'Windows'}},
        '.elf': {'function': analyze_executable, 'args': {'os_type': 'Linux'}},
        '.so': {'function': analyze_executable, 'args': {'os_type': 'Linux'}},
        '.dylib': {'function': analyze_executable, 'args': {'os_type': 'macOS'}},
        '.framework': {'function': analyze_executable, 'args': {'os_type': 'macOS'}},
        '.app': {'function': analyze_executable, 'args': {'os_type': 'macOS'}},
        '.go': {'function': analyze_golang_binary, 'args': {}},
        '.pdf': {'function': analyze_document, 'args': {}},
        '.docx': {'function': analyze_document, 'args': {}},
        '.doc': {'function': analyze_document, 'args': {}},
        '.xlsx': {'function': analyze_document, 'args': {}},
        '.pptx': {'function': analyze_document, 'args': {}},
        '.zip': {'function': analyze_archive, 'args': {}},
        '.rar': {'function': analyze_archive, 'args': {}},
        '.7z': {'function': analyze_archive, 'args': {}},
        '.tar': {'function': analyze_archive, 'args': {}},
        '.gz': {'function': analyze_archive, 'args': {}},
        '.pcap': {'function': analyze_pcap, 'args': {}},
        '.pcapng': {'function': analyze_pcap, 'args': {}},
        '.ps1': {'function': analyze_powershell, 'args': {}},
        '.eml': {'function': analyze_email, 'args': {}},
        '.msg': {'function': analyze_email, 'args': {}},
    }
    
    analysis_info = analysis_map.get(file_ext)
    
    if analysis_info:
        func = analysis_info['function']
        args_dict = analysis_info.get('args', {})
        
        console.print(f"{INFO}[>] File type detected: {file_ext}{RESET}")
        console.print(f"{INFO}[>] Using analyzer: {func.__name__}{RESET}")
        
        try:
            # Call the analysis function
            if func == analyze_executable:
                result = func(file_path, base_dir, **args_dict)
            else:
                result = func(file_path, base_dir)
            
            console.print(f"{SUCCESS}[+] Primary analysis completed{RESET}")
            
        except Exception as e:
            console.print(f"{ERROR}[-] Analysis failed: {e}{RESET}")
            logging.error(f"Analysis failed: {e}", exc_info=True)
            # Continue with pattern search if possible
    
    else:
        console.print(f"{WARNING}[!] Unsupported file type: {file_ext}{RESET}")
        console.print(f"{INFO}[>] Attempting generic analysis{RESET}")
        logging.warning(f"Unsupported file type: {file_ext}")
    
    # Perform pattern search on extracted/analyzed content
    console.print(f"\n{INFO}[>] Performing Pattern Search Analysis{RESET}")
    
    # Determine which directory to search
    search_dirs = []
    
    # Check for common extraction directories
    possible_dirs = [
        os.path.join(base_dir, "classes.dex.out"),
        os.path.join(base_dir, "extracted"),
        os.path.join(base_dir, "unzipped"),
        base_dir
    ]
    
    for dir_path in possible_dirs:
        if os.path.exists(dir_path) and os.path.isdir(dir_path):
            search_dirs.append(dir_path)
    
    if not search_dirs:
        search_dirs = [base_dir]
    
    all_results = {}
    for search_dir in search_dirs:
        console.print(f"{INFO}[>] Searching in: {search_dir}{RESET}")
        results = search_patterns(search_dir, PATTERNS)
        
        # Merge results
        for pattern, matches in results.items():
            if pattern not in all_results:
                all_results[pattern] = []
            all_results[pattern].extend(matches)
    
    # Display results
    display_results(all_results)
    
    # Run custom information extraction if any
    custom_results = {}
    if custom_patterns:
        for search_dir in search_dirs:
            custom_res = custom_info_extractor(search_dir, custom_patterns)
            for pattern, matches in custom_res.items():
                if pattern not in custom_results:
                    custom_results[pattern] = []
                custom_results[pattern].extend(matches)
        
        if custom_results:
            console.print(f"\n{INFO}[>] Custom Pattern Search Results{RESET}")
            display_results(custom_results)
    
    # Generate reports if not disabled
    if not args.no_report:
        generate_html_report(all_results, base_dir, custom_results)
        generate_comprehensive_report(all_results, custom_results, base_dir)
    
    # Perform additional analyses based on file type
    if file_ext == '.apk':
        decompiled_dir = os.path.join(base_dir, "classes.dex.out")
        if os.path.exists(decompiled_dir):
            # Additional metadata extraction
            extract_endpoints_and_urls(decompiled_dir, base_dir)
            extract_sql_information(decompiled_dir, base_dir)
            extract_hidden_domains(decompiled_dir, base_dir)
            extract_user_credentials(decompiled_dir, base_dir)
            extract_ip_ports(decompiled_dir, base_dir)
            extract_api_information(decompiled_dir, base_dir)
            extract_folder_paths(decompiled_dir, base_dir)
        
        # Advanced analyses
        analyze_uefi_injections(file_path, base_dir)
        analyze_hardware_injections(file_path, base_dir)
        analyze_memory_injections(file_path, base_dir)
        analyze_module_injections(file_path, base_dir)
        analyze_advanced_tampering(file_path, base_dir)
        analyze_firmware_manipulation(file_path, base_dir)
        analyze_process_injections(file_path, base_dir)
        analyze_firmware(file_path, base_dir)
        
        # Extract strings and hex dump
        extract_strings(file_path, base_dir)
        classes_dex_path = os.path.join(base_dir, "unzipped", "classes.dex")
        if os.path.exists(classes_dex_path):
            hex_dump(classes_dex_path, base_dir)
    
    elif file_ext in ['.exe', '.dll', '.msi', '.bin', '.elf', '.so']:
        # Advanced analyses for executables
        analyze_uefi_injections(file_path, base_dir)
        analyze_hardware_injections(file_path, base_dir)
        analyze_memory_injections(file_path, base_dir)
        analyze_module_injections(file_path, base_dir)
        analyze_advanced_tampering(file_path, base_dir)
        analyze_firmware_manipulation(file_path, base_dir)
        analyze_process_injections(file_path, base_dir)
        analyze_firmware(file_path, base_dir)
    
    # Remediation Suggestions
    suggestions = remediation_suggestions(all_results)
    if suggestions:
        console.print(f"\n{INFO}[>] Remediation Suggestions{RESET}")
        console.print("=" * 50)
        for i, suggestion in enumerate(suggestions, 1):
            console.print(f"{i}. {suggestion}")
        logging.info(f"Remediation suggestions: {suggestions}")
    
    # Summary
    console.print(f"\n{SUCCESS}[+] Analysis Completed Successfully!{RESET}")
    console.print(f"{INFO}[>] Log file: {log_file}{RESET}")
    console.print(f"{INFO}[>] Output directory: {base_dir}{RESET}")
    
    # List generated files
    console.print(f"\n{INFO}[>] Generated Files:{RESET}")
    for root, dirs, files in os.walk(base_dir):
        level = root.replace(base_dir, '').count(os.sep)
        indent = ' ' * 2 * level
        console.print(f"{indent}{os.path.basename(root)}/")
        subindent = ' ' * 2 * (level + 1)
        for file in files[:10]:  # Show first 10 files per directory
            console.print(f"{subindent}{file}")
        if len(files) > 10:
            console.print(f"{subindent}... and {len(files) - 10} more files")
    
    logging.info("Analysis Completed Successfully")

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        console.print(f"\n{WARNING}[!] Analysis interrupted by user{RESET}")
        logging.warning("Analysis interrupted by user")
        sys.exit(1)
    except Exception as e:
        console.print(f"\n{ERROR}[-] Unexpected error: {e}{RESET}")
        logging.error(f"Unexpected error: {e}", exc_info=True)
        sys.exit(1)
